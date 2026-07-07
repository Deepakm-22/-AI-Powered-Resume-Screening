import os
import docx

def create_resume_docx(file_path, name, email, phone, skills, experience, education, raw_text_block):
    """Creates a sample DOCX resume."""
    doc = docx.Document()
    
    # Title/Header
    doc.add_heading(name, level=0)
    
    # Contact Info
    doc.add_paragraph(f"Email: {email} | Phone: {phone}")
    doc.add_paragraph("Address: Hyderabad, India | LinkedIn: linkedin.com/in/candidate")
    
    # Skills Section
    doc.add_heading("Technical Skills", level=1)
    doc.add_paragraph(", ".join(skills))
    
    # Experience Section
    doc.add_heading("Professional Experience", level=1)
    doc.add_paragraph(experience)
    
    # Education Section
    doc.add_heading("Education & Credentials", level=1)
    doc.add_paragraph(education)
    
    # Full Text Block (for indexing/text extraction tests)
    doc.add_heading("Additional Information & Profile Details", level=2)
    doc.add_paragraph(raw_text_block)
    
    # Save file
    doc.save(file_path)
    print(f"Created mock resume at: {file_path}")

def main():
    # Create target directory
    output_dir = "mock_resumes"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    # Resume 1: Java Developer (Sattaru Venkata Abhilash)
    create_resume_docx(
        os.path.join(output_dir, "Sattaru_Venkata_Abhilash.docx"),
        "SATTARU VENKATA ABHILASH",
        "sattaruvenkataabhilash@gmail.com",
        "+91 9398177084",
        ["java", "spring boot", "hibernate", "microservices", "sql", "rest api", "git", "maven", "html", "css", "python", "redis", "mongodb", "ai"],
        "EXPERIENCE Intel Unnati Industrial Training – 2024 - Secure File Encryption and Decryption Application\n"
        "Duration: May 2024 - July 2024 | Role: Team Leader\n"
        "• Engineered a secure file encryption & decryption system leveraging AES -256 and Key Derivation Functions (KDF), with workflows for random File Encryption Key (FEK) generation and secure storage to safeguard sensitive data.\n"
        "• Developed an intuitive CustomTkinter GUI, simplifying complex encryption processes into a user-friendly format.\n"
        "• 3 years of overall coding and development experience in software applications.",
        "EDUCATION BTech in Computer Science and Engineering Gitam University 2021 – 2025 | Bangalore, India.\n"
        "• Secured a spot in a hackathon for a cloud-based project, highlighting innovation and problem-solving abilities.\n"
        "12th Grade Sri Chaitanya Jr College - MPC Mayuri Bhavan (2019 -2021).",
        "Full summary profile: Enthusiastic Software Engineer specialized in Java full stack development. Familiar with relational and non-relational databases like PostgreSQL, MySQL, and MongoDB. Experience with RESTful APIs, Git, Maven, and building secure cryptographic modules. Goal-oriented team worker."
    )
    
    # Resume 2: Python / Django Developer (Ajay Kumar)
    create_resume_docx(
        os.path.join(output_dir, "Ajay_Kumar.docx"),
        "Ajay Kumar",
        "ajayoneness123@gmail.com",
        "+91 9876543210",
        ["python", "django", "machine learning", "nlp", "rest api", "tensorflow", "scikit-learn", "nltk", "git", "html", "css", "sqlite"],
        "EXPERIENCE Senior Backend Developer - Python/Django Solutions\n"
        "Duration: May 2022 - Present | Role: Full Stack Python Developer\n"
        "• Designed and engineered web applications using Django and Django REST Framework, resulting in a 40% speedup of data APIs.\n"
        "• Integrated machine learning models for semantic searches and data categorization using scikit-learn and spaCy NLP frameworks.\n"
        "• Formulate automated scripts for automated text parsing, saving 20+ manual team-hours per week.\n"
        "• 4+ years of professional backend programming experience.",
        "EDUCATION Bachelor of Technology in Computer Science and Engineering\n"
        "IIT Delhi (2018 - 2022) | GPA: 8.9/10\n"
        "• Specialized in Artificial Intelligence and Machine Learning courses.",
        "Profile summary: Highly skilled Python and Django backend developer with interest in Natural Language Processing and Machine Learning models. Adept at developing clean APIs, parsing raw data structures, and deploying scalable software architectures with Docker, Git, and Linux environments."
    )
    
    # Resume 3: Marketing Executive (John Doe)
    create_resume_docx(
        os.path.join(output_dir, "John_Doe.docx"),
        "John Doe",
        "johndoe@example.com",
        "+1 555 0199",
        ["marketing", "sales", "social media", "negotiation", "public speaking", "advertising", "team management"],
        "EXPERIENCE Sales and Marketing Coordinator - Brand Dynamics Corp\n"
        "Duration: June 2020 - January 2025 | Role: Lead Specialist\n"
        "• Supervised social media campaign budgets, achieving 15% increase in conversion rates.\n"
        "• Directed public relations efforts, pitching brand strategies to 50+ corporate clients.\n"
        "• Negotiated commercial contracts, expanding regional business growth by $100K.\n"
        "• 5 years of experience in product branding, sales, and retail marketing.",
        "EDUCATION Bachelor of Business Administration (BBA) in Marketing\n"
        "Boston University (2016 - 2020)",
        "Summary: Dynamic sales representative and branding expert. Skilled at managing cross-functional teams, executing high-ROI social media ads, negotiating commercial deals, and managing public relations accounts."
    )

if __name__ == '__main__':
    main()
