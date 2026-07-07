import re
import os
import PyPDF2
import docx
import spacy

# Load spaCy model for Named Entity Recognition
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    # Fallback if model not downloaded yet (will download in setup)
    nlp = None

# A comprehensive dictionary of skills to scan for (common CSE/IT skills)
SKILLS_DICTIONARY = [
    # Programming Languages
    "python", "java", "c++", "c#", "javascript", "typescript", "ruby", "php", "go", "rust", "kotlin", "swift", "scala", "r", "sql", "html", "css",
    # Frameworks & Libraries
    "django", "flask", "fastapi", "spring boot", "spring", "hibernate", "microservices", "react", "angular", "vue", "next.js", "node.js", "express",
    "bootstrap", "tailwind", "jquery", "scikit-learn", "tensorflow", "pytorch", "keras", "spacy", "nltk", "opencv", "pandas", "numpy",
    # Databases
    "mysql", "postgresql", "sqlite", "mongodb", "redis", "oracle", "sql server", "cassandra", "dynamodb",
    # Tools, Platforms & DevOps
    "git", "github", "docker", "kubernetes", "aws", "azure", "gcp", "heroku", "jenkins", "ansible", "terraform", "linux", "unix",
    # Methodologies & Concepts
    "machine learning", "deep learning", "nlp", "natural language processing", "computer vision", "artificial intelligence",
    "rest api", "restful api", "graphql", "agile", "scrum", "data science", "oop", "object oriented programming"
]

def extract_text_from_pdf(file_path):
    """Extracts raw text from a PDF file using PyPDF2."""
    text = ""
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return text

def extract_text_from_docx(file_path):
    """Extracts raw text from a DOCX file using python-docx."""
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return text

def extract_text(file_path):
    """Wrapper to extract text based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    else:
        return ""

def clean_text(text):
    """Preprocesses text for parsing."""
    # Replace multiple spaces/newlines with single ones
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def extract_email(text):
    """Extracts email address using Regex."""
    pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
    match = re.search(pattern, text)
    return match.group(0) if match else "Not Found"

def extract_phone(text):
    """Extracts phone numbers using Regex."""
    # Match standard formats like +91 9876543210, 9876543210, +1-234-567-8901, etc.
    pattern = r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}|\+?\d[\d\s.-]{9,13}\d'
    match = re.search(pattern, text)
    return match.group(0).strip() if match else "Not Found"

def extract_name(text):
    """Extracts candidate name using spaCy NER or fallback rules."""
    if nlp:
        doc = nlp(text[:1000])  # Scan only first 1000 chars for speed and name location
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                name = ent.text.strip()
                # Ensure it's not a common technical keyword or stopword and has letters/spaces
                if len(name.split()) >= 2 and len(name.split()) <= 4:
                    if not any(kw in name.lower() for kw in ['curriculum', 'resume', 'cv', 'experience', 'email', 'phone', 'apply', 'job', 'work']):
                        # Only allow letters, spaces, dots, and hyphens
                        if re.match(r'^[a-zA-Z\s\.\-]+$', name):
                            return name
    
    # Fallback: Extract first line of non-empty text, assuming it is the name
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    for line in lines[:5]:  # Look at the first 5 lines
        # Clean up line if it contains common resume headers
        if not any(header in line.lower() for header in ['resume', 'curriculum', 'cv', 'profile', 'contact', 'email', 'phone']):
            # Only allow letters, spaces, dots, and hyphens
            if re.match(r'^[a-zA-Z\s\.\-]+$', line):
                words = line.split()
                if len(words) >= 2 and len(words) <= 4:
                    return line
    
    return "Candidate Name"

def extract_skills(text):
    """Scans text for skills listed in the technical skills dictionary."""
    extracted = []
    text_lower = text.lower()
    
    # Simple word/phrase boundary matching
    for skill in SKILLS_DICTIONARY:
        # Avoid partial matches, e.g. "git" matching in "digital"
        # We use regex with word boundaries. Special handling for skills with symbols like c++, c#
        escaped_skill = re.escape(skill)
        # Handle special characters at start/end of skills
        pattern = rf'\b{escaped_skill}\b'
        if skill == 'c++':
            pattern = r'\bc\+\+'
        elif skill == 'c#':
            pattern = r'\bc\#'
        elif skill == '.net':
            pattern = r'\.net\b'
            
        if re.search(pattern, text_lower):
            extracted.append(skill)
            
    return list(set(extracted))

def extract_section(text, keywords):
    """Heuristic extraction of a section based on heading keywords."""
    lines = text.split('\n')
    section_content = []
    in_section = False
    
    # Clean keywords list
    keywords_lower = [kw.lower() for kw in keywords]
    
    # A list of headers that would signify the end of a section
    all_headers = ["experience", "employment", "work history", "education", "skills", "projects", "certifications", "interests", "languages", "profile", "summary"]
    
    for line in lines:
        line_clean = line.strip().lower()
        if not line_clean:
            continue
            
        # Check if line indicates start of target section
        if any(kw in line_clean for kw in keywords_lower) and len(line_clean) < 30:
            in_section = True
            section_content.append(line)
            continue
            
        # Check if line indicates start of a DIFFERENT section (exit condition)
        if in_section:
            if any(h in line_clean for h in all_headers if h not in keywords_lower) and len(line_clean) < 30:
                in_section = False
                break
            section_content.append(line)
            
    # Fallback: if we didn't find section by structure, grab sentences containing the keywords
    if not section_content:
        for line in lines:
            if any(kw in line.lower() for kw in keywords_lower):
                section_content.append(line)
                
    return "\n".join(section_content[:15])  # Cap to first 15 lines of content to avoid giant text blocks

def parse_resume(file_path):
    """Main parsing interface: Extracts raw text and parses all properties."""
    raw_text = extract_text(file_path)
    clean_text_str = clean_text(raw_text)
    
    name = extract_name(raw_text)
    email = extract_email(clean_text_str)
    phone = extract_phone(clean_text_str)
    skills = extract_skills(clean_text_str)
    
    experience_text = extract_section(raw_text, ["experience", "work history", "employment", "professional background"])
    education_text = extract_section(raw_text, ["education", "academic qualification", "university", "school", "degree"])
    
    # Fallbacks if section parsing is completely blank
    if not experience_text.strip():
        experience_text = "Experience details not explicitly found. Raw text contains background info."
    if not education_text.strip():
        education_text = "Education details not explicitly found. Raw text contains academic info."
        
    return {
        'name': name,
        'email': email,
        'phone': phone,
        'skills': skills,
        'experience_text': experience_text,
        'education_text': education_text,
        'raw_text': raw_text
    }
